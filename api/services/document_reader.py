import io
import logging

import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text content from a .docx file, including tables."""
    doc = Document(io.BytesIO(file_bytes))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.replace("|", "").strip():
                parts.append(row_text)

    return "\n".join(parts)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text content from a PDF file, including tables."""
    parts: list[str] = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                parts.append(text)

            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_text = " | ".join(
                        (cell or "").strip() for cell in row
                    )
                    if row_text.replace("|", "").strip():
                        parts.append(row_text)

    return "\n".join(parts)


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Extract text from a file based on its extension."""
    lower = filename.lower()

    if lower.endswith(".docx"):
        logger.info("Extracting text from .docx: %s", filename)
        return extract_text_from_docx(file_bytes)
    elif lower.endswith(".pdf"):
        logger.info("Extracting text from PDF: %s", filename)
        return extract_text_from_pdf(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {filename}")
