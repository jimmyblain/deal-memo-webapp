"""
Run this script once to generate the deal_memo_template.docx file.
Requires: pip install python-docx docxtpl

Usage: python templates/create_template.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading_style(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0x78, 0xD4)
    return heading


def add_field_row(table, label: str, placeholder: str):
    row = table.add_row()
    label_cell = row.cells[0]
    value_cell = row.cells[1]

    label_cell.text = label
    for p in label_cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

    value_cell.text = placeholder
    for p in value_cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(10)


def main():
    doc = Document()

    # Title
    title = doc.add_heading("DEAL MEMO", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0x78, 0xD4)

    doc.add_paragraph("")  # spacer

    # Deal Overview section
    add_heading_style(doc, "Deal Overview", level=1)
    table1 = doc.add_table(rows=0, cols=2)
    table1.style = "Table Grid"
    table1.columns[0].width = Inches(2.0)
    table1.columns[1].width = Inches(4.5)

    add_field_row(table1, "Vendor Name", "{{ vendor_name }}")
    add_field_row(table1, "Deal Owner", "{{ deal_owner }}")
    add_field_row(table1, "Department", "{{ department }}")
    add_field_row(table1, "Priority", "{{ deal_priority }}")
    add_field_row(table1, "Budget Code", "{{ budget_code }}")

    doc.add_paragraph("")

    # Service Details section
    add_heading_style(doc, "Service Details", level=1)
    table2 = doc.add_table(rows=0, cols=2)
    table2.style = "Table Grid"
    table2.columns[0].width = Inches(2.0)
    table2.columns[1].width = Inches(4.5)

    add_field_row(table2, "Description of Services", "{{ description_of_services }}")
    add_field_row(table2, "Key Deliverables", "{{ key_deliverables }}")
    add_field_row(table2, "Contract Type", "{{ contract_type }}")
    add_field_row(table2, "SLA Terms", "{{ sla_terms }}")

    doc.add_paragraph("")

    # Financial Terms section
    add_heading_style(doc, "Financial Terms", level=1)
    table3 = doc.add_table(rows=0, cols=2)
    table3.style = "Table Grid"
    table3.columns[0].width = Inches(2.0)
    table3.columns[1].width = Inches(4.5)

    add_field_row(table3, "Total Cost", "{{ total_cost }}")
    add_field_row(table3, "Payment Terms", "{{ payment_terms }}")
    add_field_row(table3, "Liability Cap", "{{ liability_cap }}")
    add_field_row(table3, "Insurance Requirements", "{{ insurance_requirements }}")

    doc.add_paragraph("")

    # Contract Terms section
    add_heading_style(doc, "Contract Terms", level=1)
    table4 = doc.add_table(rows=0, cols=2)
    table4.style = "Table Grid"
    table4.columns[0].width = Inches(2.0)
    table4.columns[1].width = Inches(4.5)

    add_field_row(table4, "Start Date", "{{ contract_start_date }}")
    add_field_row(table4, "End Date", "{{ contract_end_date }}")
    add_field_row(table4, "Renewal Terms", "{{ renewal_terms }}")
    add_field_row(table4, "Termination Clause", "{{ termination_clause }}")
    add_field_row(table4, "Confidentiality Terms", "{{ confidentiality_terms }}")

    doc.add_paragraph("")

    # Business Justification section
    add_heading_style(doc, "Business Justification", level=1)
    doc.add_paragraph("{{ business_justification }}")

    doc.add_paragraph("")

    # Approval section
    add_heading_style(doc, "Approval", level=1)
    table5 = doc.add_table(rows=0, cols=2)
    table5.style = "Table Grid"
    table5.columns[0].width = Inches(2.0)
    table5.columns[1].width = Inches(4.5)

    add_field_row(table5, "Approver", "{{ approver_name }}")
    add_field_row(table5, "Signature", "")
    add_field_row(table5, "Date", "")

    doc.add_paragraph("")

    # Internal Notes section
    add_heading_style(doc, "Internal Notes", level=1)
    doc.add_paragraph("{{ internal_notes }}")

    # Save
    output_path = "templates/deal_memo_template.docx"
    doc.save(output_path)
    print(f"Template created: {output_path}")


if __name__ == "__main__":
    main()
